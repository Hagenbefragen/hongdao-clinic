import os
import sys
from datetime import datetime
from difflib import SequenceMatcher
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_run(p, text, is_ins=False, is_del=False, author="Antigravity", date_str=None):
    if not date_str:
        date_str = datetime.utcnow().strftime('%Y-%m-%dT%H:%M:%SZ')
        
    if is_ins:
        # <w:ins w:id="..." w:author="..." w:date="...">
        ins = OxmlElement('w:ins')
        ins.set(qn('w:id'), str(id(ins)))
        ins.set(qn('w:author'), author)
        ins.set(qn('w:date'), date_str)
        
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        
        r.append(t)
        ins.append(r)
        p._element.append(ins)
    elif is_del:
        # <w:del w:id="..." w:author="..." w:date="...">
        delete = OxmlElement('w:del')
        delete.set(qn('w:id'), str(id(delete)))
        delete.set(qn('w:author'), author)
        delete.set(qn('w:date'), date_str)
        
        r = OxmlElement('w:r')
        delText = OxmlElement('w:delText')
        delText.text = text
        delText.set(qn('xml:space'), 'preserve')
        
        r.append(delText)
        delete.append(r)
        p._element.append(delete)
    else:
        # standard run
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p._element.append(r)

def generate_redline_paragraph(p, original_text, new_text):
    # clear existing elements
    for child in list(p._element):
        p._element.remove(child)
        
    matcher = SequenceMatcher(None, original_text, new_text)
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            add_run(p, original_text[i1:i2], is_ins=False, is_del=False)
        elif tag == 'replace':
            add_run(p, original_text[i1:i2], is_ins=False, is_del=True)
            add_run(p, new_text[j1:j2], is_ins=True, is_del=False)
        elif tag == 'delete':
            add_run(p, original_text[i1:i2], is_ins=False, is_del=True)
        elif tag == 'insert':
            add_run(p, new_text[j1:j2], is_ins=True, is_del=False)

def apply_redlines(original_docx_path, output_docx_path, mapped_changes):
    doc = Document(original_docx_path)
    
    # Enable track changes globally
    settings = doc.settings.element
    if settings.find(qn('w:trackRevisions')) is None:
        track_revs = OxmlElement('w:trackRevisions')
        settings.append(track_revs)
        
    for i, p in enumerate(doc.paragraphs):
        if i in mapped_changes:
            new_text = mapped_changes[i]
            original_text = p.text
            generate_redline_paragraph(p, original_text, new_text)
            
    # Also add the new paragraphs at the end if we have extra content
    max_idx = max(mapped_changes.keys()) if mapped_changes else -1
    for i in sorted(mapped_changes.keys()):
        if i >= len(doc.paragraphs):
            new_p = doc.add_paragraph()
            generate_redline_paragraph(new_p, "", mapped_changes[i])
            
    doc.save(output_docx_path)
    print(f"Redlined document saved to {output_docx_path}")

if __name__ == "__main__":
    original = r'c:\ohm\Documentation\VC-EXIT\target-pitches\Fortress-Pitch\LOI-FANCCI-Destill.AI.docx'
    output = r'c:\ohm\Documentation\VC-EXIT\target-pitches\Fortress-Pitch\LOI-FANCCI-Destill.AI-Redlined.docx'
    
    changes = {
        6: "1. Hintergrund und Zielsetzung\nZiel dieses Letters of Intent ist es, die Grundlage für eine strategische Zusammenarbeit zwischen FANCCI und DESTILLE zu definieren. FANCCI entwickelt derzeit ihre Creator Plattform („FANCCI 2.0“), mit Fokus auf Monetarisierung und Community-basierte Einnahmemodelle.",
        7: "DESTILLE verfügt über eine proprietäre Softwarelösung („Destille AI“), die als Compliance-Shield gegen die regulatorischen Anforderungen der EU Digital Services Act (DSA) (Umsatz-Sanktionen bis zu 6% bei Urheberrechtsverletzungen) und die CDSM-Richtlinie (Art. 17 – Staydown-Deadline Juni 2026) fungiert. Die Technologie ermöglicht den Schutz digitaler Inhalte durch generative Wasserzeichen, globales Monitoring und automatisierte rechtliche Durchsetzung.",
        8: "Nach Angaben von DESTILLE basiert die Technologie auf einem weltweit führenden Portfolio von über 2.200 Patentansprüchen (Claims), insbesondere in den USA.",
        11: "2. Gegenstand der Zusammenarbeit",
        12: "Die Parteien beabsichtigen die Integration der DESTILLE-Technologie in die FANCCI-Plattform als zentrales Schutzsystem („Regulatory-Shield-as-a-Service“). Dies umfasst:",
        14: " a. Content Protection & Provenance",
        15: "• Integration unsichtbarer Wasserzeichen (Staging/Production).",
        16: " b. Monitoring & Global Detection",
        17: " • Tracking von Piracy-Bots und unautorisierten Netzwerken.",
        18: " c. Automated Enforcement",
        19: " • Bereitstellung gerichtsverwertbarer Berichte und Unterstützung bei Takedowns/Litigation.",
        21: "3. Geplante Integration",
        22: "Die Parteien beabsichtigen eine schrittweise Einführung (Pilotphase 12 Wochen → Rollout). DESTILLE stellt hierfür die erforderlichen Schnittstellen (API/SDK) bereit.",
        23: "",
        25: "4. IP-Souveränität und Strategische Partnerschaft",
        26: "Die Parteien vereinbaren folgendes für die Sicherung der Wettbewerbsvorteile bei gleichzeitiger Wahrung der unternehmerischen Souveränität von DESTILLE:",
        27: "• Launch-Exklusivität: DESTILLE gewährt FANCCI für einen Zeitraum von __ Monaten ab Unterzeichnung eine exklusive Nutzung der Technologie innerhalb der spezifischen Markt-Nische \"Creator & Influencer Hub Platforms\".",
        28: "• IP Ownership & Transferability: Sämtliche Rechte am geistigen Eigentum (IP) verbleiben ausschließlich bei DESTILLE. DESTILLE behält sich ausdrücklich das Recht vor, das IP-Portfolio jederzeit an Dritte zu veräußern oder zu übertragen (Exit/M&A).",
        29: "• Succession-Guarantee: Im Falle einer Veräußerung an einen Dritten (z. B. Big-Tech-Unternehmenskauf) verpflichtet sich DESTILLE, sicherzustellen, dass die zum Zeitpunkt des Verkaufs bestehende operative Implementierung und Lizenzierung für FANCCI durch den Rechtsnachfolger fortgeführt wird (\"Right to use remains\").",
        31: "5. Wirtschaftliche Rahmenbedingungen",
        32: "Die Details werden in einem separaten Vertrag geregelt. Mögliche Modelle: SaaS-Lizenz, Revenue-Share auf durchgesetzte Forderungen (Enforcement-Success) oder Kombinationen daraus.",
        33: "",
        35: "6. Vertraulichkeit (Binding)",
        36: "Die Parteien behandeln alle Informationen (technisch/geschäftlich) streng vertraulich. Dies gilt insbesondere für Details zum Patent-Portfolio und strategische Roadmap-Pläne.",
        38: "7. Rechtliche Unverbindlichkeit",
        39: "Dieser LOI dokumentiert die aktuelle Absicht der Parteien.",
        42: "Rechtlich bindend sind ausschließlich die Bestimmungen zur Vertraulichkeit (Ziff. 6) sowie die exklusive Verhandlungsphase (Ziff. 4). Ein verbindlicher Vertrag wird innerhalb von 12 Wochen finalisiert.",
        43: "",
        44: "",
        45: "",
        51: "8. Anwendbares Recht",
        89: "Für FANCCI Group AG"
    }
    
    apply_redlines(original, output, changes)
